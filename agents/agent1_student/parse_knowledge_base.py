#!/usr/bin/env python3
"""
Script to parse knowledge base files and convert them to Qdrant-compatible format.
Supports: txt, docx, doc, pdf files
"""

import os
import json
import re
from pathlib import Path
from typing import List, Dict, Any
import uuid

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. DOCX support will be limited.")

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("Warning: PyPDF2 not installed. PDF support will be limited.")

try:
    import python_docx
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


class KnowledgeBaseParser:
    def __init__(self, source_dir: str, output_dir: str):
        self.source_dir = Path(source_dir)
        self.output_dir = Path(output_dir)
        self.documents = []
        
    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from .txt file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT file {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from .docx file"""
        if not HAS_DOCX:
            return ""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error reading DOCX file {file_path}: {e}")
            return ""
    
    def extract_text_from_doc(self, file_path: Path) -> str:
        """Extract text from .doc file (legacy format)"""
        try:
            # Try using python-docx for newer .doc files
            if HAS_DOCX:
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return text
        except Exception as e:
            print(f"Note: Could not parse .doc file {file_path} with python-docx: {e}")
        
        # Fallback: try to extract text using other methods
        return f"[Document content not fully extractable: {file_path.name}]"
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from .pdf file"""
        if not HAS_PDF:
            return f"[PDF content not extractable - PyPDF2 not installed: {file_path.name}]"
        try:
            text = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            return "\n".join(text)
        except Exception as e:
            print(f"Error reading PDF file {file_path}: {e}")
            return f"[PDF content not fully extractable: {file_path.name}]"
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from any supported file format"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.txt':
            return self.extract_text_from_txt(file_path)
        elif suffix == '.docx':
            return self.extract_text_from_docx(file_path)
        elif suffix == '.doc':
            return self.extract_text_from_doc(file_path)
        elif suffix == '.pdf':
            return self.extract_text_from_pdf(file_path)
        else:
            return ""
    
    def create_chunks(self, text: str, chunk_size: int = 500) -> List[str]:
        """Split text into chunks for better vectorization"""
        if not text:
            return []
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return [c for c in chunks if c]
    
    def parse_files(self):
        """Parse all files in the source directory structure"""
        for category_dir in self.source_dir.iterdir():
            if not category_dir.is_dir():
                continue
            
            category = category_dir.name
            print(f"Processing category: {category}")
            
            for file_path in category_dir.glob('*'):
                if file_path.is_file() and file_path.suffix.lower() in ['.txt', '.docx', '.doc', '.pdf']:
                    print(f"  - Parsing: {file_path.name}")
                    
                    text = self.extract_text(file_path)
                    if not text:
                        continue
                    
                    chunks = self.create_chunks(text)
                    
                    for chunk in chunks:
                        doc = {
                            "id": str(uuid.uuid4()),
                            "category": category,
                            "source_file": file_path.name,
                            "content": chunk,
                            "metadata": {
                                "category": category,
                                "source": file_path.name,
                                "file_type": file_path.suffix.lower()
                            }
                        }
                        self.documents.append(doc)
    
    def create_output_structure(self):
        """Create the output directory structure"""
        categories = set([doc["category"] for doc in self.documents])
        for category in categories:
            category_dir = self.output_dir / category
            category_dir.mkdir(parents=True, exist_ok=True)
            
            # Save documents for this category
            category_docs = [doc for doc in self.documents if doc["category"] == category]
            
            # Save as JSON
            with open(category_dir / f"{category}_documents.json", 'w', encoding='utf-8') as f:
                json.dump(category_docs, f, ensure_ascii=False, indent=2)
            
            print(f"Created: {category_dir}/{category}_documents.json ({len(category_docs)} chunks)")
    
    def save_all_documents(self):
        """Save all documents as a single JSON file"""
        all_docs_path = self.output_dir / "all_documents.json"
        with open(all_docs_path, 'w', encoding='utf-8') as f:
            json.dump(self.documents, f, ensure_ascii=False, indent=2)
        print(f"\nTotal documents created: {len(self.documents)}")
        print(f"Saved to: {all_docs_path}")


def main():
    source_dir = "/home/admin/Desktop/ai_stack/ai_stack/agents/agent1_student/chatbot-baza-wiedzy-nowa"
    output_dir = "/home/admin/Desktop/ai_stack/ai_stack/agents/agent1_student/knowledge"
    
    parser = KnowledgeBaseParser(source_dir, output_dir)
    print("Starting knowledge base parsing...\n")
    parser.parse_files()
    parser.create_output_structure()
    parser.save_all_documents()
    print("\nParsing complete!")


if __name__ == "__main__":
    main()
